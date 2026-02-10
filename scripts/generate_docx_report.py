#!/usr/bin/env python3
"""
Generate TTG Supervisor Report DOCX for Milestone 3.

Uses the existing M1+M2 report as a template and appends M3 data,
keeping the same document styles and formatting.

Usage:
    python3 scripts/generate_docx_report.py
"""

import copy
import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx")
    sys.exit(1)


PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(
    os.path.dirname(PROJECT_DIR),  # Go up from worktree
    "docs", "tracking", "TTG_Project_Status_Report_M1_M2.docx"
)
# Fallback: look in main repo
if not os.path.exists(TEMPLATE_PATH):
    TEMPLATE_PATH = os.path.join(
        "/home/xavierand_/Desktop/TTG",
        "docs", "tracking", "TTG_Project_Status_Report_M1_M2.docx"
    )

OUTPUT_PATH = os.path.join(
    PROJECT_DIR, "docs", "tracking", "TTG_Project_Status_Report_M1_M2_M3.docx"
)


def add_table_row(table, cells_data, bold_first=False):
    """Add a row to an existing table."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(text)
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row


def _get_style(doc, name):
    """Get a style by name, falling back to iterating all styles."""
    try:
        return doc.styles[name]
    except KeyError:
        # Pandoc-generated headings can't be accessed by name;
        # iterate and match manually.
        for s in doc.styles:
            if s.name == name:
                return s
        raise KeyError(f"Style '{name}' not found in document")


def add_heading(doc, text, level=2):
    """Add a heading using paragraph style (compatible with Pandoc templates)."""
    p = doc.add_paragraph()
    p.style = _get_style(doc, f"Heading {level}")
    p.text = text
    return p


def add_compact(doc, text):
    """Add a bullet-like compact item."""
    p = doc.add_paragraph()
    p.style = _get_style(doc, "Compact")
    p.text = text
    return p


def add_body(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph()
    p.style = _get_style(doc, "Body Text")
    p.text = text
    return p


def add_first(doc, text):
    """Add first paragraph (intro text after heading)."""
    p = doc.add_paragraph()
    p.style = _get_style(doc, "First Paragraph")
    p.text = text
    return p


def add_source(doc, text):
    """Add source code block."""
    p = doc.add_paragraph()
    p.style = _get_style(doc, "Source Code")
    p.text = text
    return p


def create_styled_table(doc, headers, rows_data, style=None):
    """Create a table with header row and data rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    if style:
        table.style = style
    else:
        # Use the 'Table' style from the template
        table.style = doc.styles["Table"]

    # Header row
    hdr = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for row_data in rows_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            row.cells[i].text = str(text)

    return table


def update_milestone_table(doc):
    """Find the milestone overview table (Table 0) and add M3 row."""
    if len(doc.tables) < 1:
        return
    table = doc.tables[0]

    # Update the existing M3 row (row 3) or add one
    # Table 0 has: Milestone, Deliverables, Status, Completion Date, Key Metrics
    found_m3 = False
    for row in table.rows:
        if "M3" in row.cells[0].text or "Milestone 3" in row.cells[0].text:
            row.cells[0].text = "M3: Production Queue"
            row.cells[1].text = "RabbitMQ backend, dual-mode demo, retry/DLQ, visual monitoring"
            row.cells[2].text = "✅ Complete"
            row.cells[3].text = "2026-02-09"
            row.cells[4].text = "100/100 chunks, 25 p/s (RabbitMQ), 27 p/s (Redis)"
            found_m3 = True
            break

    if not found_m3:
        row = table.add_row()
        row.cells[0].text = "M3: Production Queue"
        row.cells[1].text = "RabbitMQ backend, dual-mode demo, retry/DLQ, visual monitoring"
        row.cells[2].text = "✅ Complete"
        row.cells[3].text = "2026-02-09"
        row.cells[4].text = "100/100 chunks, 25 p/s (RabbitMQ), 27 p/s (Redis)"


def update_comparison_table(doc):
    """The backend comparison is already added as a new section; skip modifying old table."""
    pass  # New 3-way comparison added in add_backend_comparison()


def add_m3_section(doc):
    """Add the complete M3 section to the document."""

    # ── Milestone 3 Heading ──
    add_heading(doc, "Milestone 3: Production Queue (COMPLETE)", level=2)
    add_first(doc, "Completion Date: February 9, 2026")

    # ── Deliverables ──
    add_heading(doc, "Deliverables", level=3)

    deliverables = [
        "✅ RabbitMQ backend implementation (src/rabbitmq_queue.py)",
        "✅ Dual-backend worker toggle (QUEUE_BACKEND=redis|rabbitmq)",
        "✅ RabbitMQ Kubernetes manifests (Pod + PVC + Service)",
        "✅ Retry queue with configurable TTL + Dead Letter Queue",
        "✅ Unified one-command demo script (Redis and RabbitMQ)",
        "✅ RabbitMQ Management UI monitoring (port 15672)",
        "✅ CLI queue monitor script (rabbitmq_monitor.sh)",
        "✅ Configurable fault simulation (SIMULATE_FAULT_RATE)",
        "✅ Strict TTG-only cleanup for shared developer machines",
        "✅ Comprehensive documentation and test results",
    ]
    for d in deliverables:
        add_compact(doc, d)

    # ── Test Results ──
    add_heading(doc, "Test Results", level=3)

    add_body(doc,
        "Key Finding: Both Redis and RabbitMQ backends achieve 100% task completion "
        "with zero data loss, even when a worker is forcefully killed at ~30% progress. "
        "RabbitMQ automatically requeues in-flight messages on worker disconnect. "
        "Redis uses XCLAIM-based stale task recovery. "
        "Throughput is comparable: RabbitMQ at 25 params/sec vs Redis at 27 params/sec "
        "(~7% difference due to AMQP protocol overhead)."
    )

    # Fault tolerance results table
    create_styled_table(doc,
        ["Metric", "Result"],
        [
            ["Total Chunks", "100/100 completed (both backends)"],
            ["Workers Deployed", "3 (standalone pods)"],
            ["Worker Killed At", "~30-36% progress"],
            ["Total Time (no fault)", "39s (RabbitMQ), 36s (Redis)"],
            ["Total Time (with fault)", "49s (RabbitMQ), 48s (Redis)"],
            ["Throughput (no fault)", "25 p/s (RabbitMQ), 27 p/s (Redis)"],
            ["Data Loss", "ZERO — all tasks completed"],
        ],
    )

    # ── RabbitMQ Queue Stats ──
    add_heading(doc, "RabbitMQ Queue Stats (After Completion)", level=3)

    create_styled_table(doc,
        ["Queue", "Messages", "Status"],
        [
            ["ttg.tasks", "0", "All consumed"],
            ["ttg.results", "100", "All results collected"],
            ["ttg.tasks.retry", "0", "No retries needed"],
            ["ttg.tasks.dlq", "0", "No dead letters"],
        ],
    )

    # ── Architecture (M3) ──
    add_heading(doc, "Architecture (M3)", level=3)

    arch_items = [
        "RabbitMQ AMQP 0-9-1 for task distribution (direct exchange → queue binding)",
        "Retry path: failed task → retry exchange → retry queue (TTL) → back to main queue",
        "Dead Letter Queue (DLQ) for tasks exceeding max retries (default: 3)",
        "Worker uses basic_get with manual acknowledgment (auto_ack=False)",
        "Native fault tolerance: unacked messages auto-requeue on consumer disconnect",
        "Dual-backend toggle: same worker code, different queue backends",
    ]
    for item in arch_items:
        add_compact(doc, item)

    doc.add_paragraph()  # spacer


def add_backend_comparison(doc):
    """Add backend comparison section."""
    add_heading(doc, "Backend Comparison: Redis vs RabbitMQ", level=2)

    create_styled_table(doc,
        ["Metric", "Redis (M2)", "RabbitMQ (M3)"],
        [
            ["Protocol", "Redis Streams (RESP)", "AMQP 0-9-1"],
            ["Task Distribution", "XREADGROUP (consumer group)", "basic_get (prefetch=1)"],
            ["Fault Recovery", "XCLAIM stale tasks", "Auto-requeue on disconnect"],
            ["Retry Logic", "Manual (stale check interval)", "Retry queue with TTL + DLQ"],
            ["Monitoring", "RedisInsight + CLI", "Management UI + CLI"],
            ["Throughput (3 workers)", "27 params/sec", "25 params/sec"],
            ["Fault Tolerance", "✅ 100% recovery", "✅ 100% recovery"],
            ["Operational Maturity", "Simpler, fewer moving parts", "Richer semantics (retry, DLQ, UI)"],
        ],
    )

    doc.add_paragraph()  # spacer


def add_m3_files(doc):
    """Add M3 key files section."""
    add_heading(doc, "Key Files Added (M3)", level=3)

    create_styled_table(doc,
        ["File", "Purpose"],
        [
            ["src/rabbitmq_queue.py", "RabbitMQ queue backend (topology, retry, DLQ)"],
            ["k8s/manifests/rabbitmq.yaml", "RabbitMQ Pod + PVC + Service"],
            ["k8s/manifests/parallel-jobs-queue-rabbitmq.yaml", "RabbitMQ worker Job manifest"],
            ["scripts/rabbitmq_monitor.sh", "CLI queue monitor (--watch mode)"],
        ],
    )

    add_heading(doc, "Updated Files (M3)", level=3)

    create_styled_table(doc,
        ["File", "Change"],
        [
            ["src/worker.py", "Dual-backend support, SIMULATE_FAULT_RATE"],
            ["docker/Dockerfile", "RabbitMQ env vars, fault simulation"],
            ["scripts/run-demo.sh", "Unified --backend redis|rabbitmq"],
            ["scripts/cleanup-ttg.sh", "Strict TTG-only, --rabbitmq flag"],
            ["docs/guides/QUEUE_MODE_GUIDE.md", "Updated for M3 (v1.3.0)"],
        ],
    )

    doc.add_paragraph()  # spacer


def update_next_steps(doc):
    """Replace the old 'Next Steps: M3 Planning' with M4."""
    # Find and update the Next Steps heading
    for i, para in enumerate(doc.paragraphs):
        if "Next Steps" in para.text and "Milestone 3" in para.text:
            para.text = "Next Steps: Milestone 4 Planning"
        # Also update old "Success Criteria for M3" to "M4"
        if "Success Criteria for M3" in para.text:
            para.text = "Success Criteria for M4"

    # Find the target dates and update
    for i, para in enumerate(doc.paragraphs):
        if "Target Start: February 10" in para.text:
            para.text = (
                "Target Start: TBD\n"
                "Target Completion: TBD\n"
            )
            break

    # Update success criteria
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Compact":
            if "☐ AKS cluster" in para.text:
                para.text = "☐ Prometheus + Grafana monitoring dashboards in Kind"
            elif "☐ Workers processing tasks in Azure" in para.text:
                para.text = "☐ Full RabbitMQ cutover (deprecate Redis path)"
            elif "☐ Monitoring dashboard operational" in para.text:
                para.text = "☐ 100K+ parameter scale test"
            elif "☐ CI/CD pipeline" in para.text:
                para.text = "☐ Azure AKS deployment (production)"
            elif "☐ 100K+ parameter scale" in para.text:
                para.text = "☐ CI/CD pipeline for automated deployment"


def update_demo_commands(doc):
    """Update the demo commands section."""
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Source Code":
            if "run-demo.sh" in para.text:
                para.text = (
                    "# Navigate to project\n"
                    "cd /home/xavierand_/Desktop/TTG\n"
                    "\n"
                    "# Run RabbitMQ demo with fault injection (RECOMMENDED)\n"
                    "./scripts/run-demo.sh --backend rabbitmq --scale small --fault-demo --monitor both\n"
                    "\n"
                    "# Run Redis demo for comparison\n"
                    "./scripts/run-demo.sh --backend redis --scale small --fault-demo --monitor both\n"
                    "\n"
                    "# Safe cleanup (TTG resources only)\n"
                    "./scripts/cleanup-ttg.sh --all --force\n"
                    "\n"
                    "# Preview cleanup (dry-run, no changes)\n"
                    "./scripts/cleanup-ttg.sh --all --dry-run"
                )
                break


def update_report_metadata(doc):
    """Update report date and version."""
    for para in doc.paragraphs:
        if "Report Date: February 3, 2026" in para.text:
            para.text = para.text.replace(
                "Report Date: February 3, 2026",
                "Report Date: February 9, 2026"
            )
        if "Report Version: 2.0" in para.text:
            para.text = para.text.replace(
                "Report Version: 2.0",
                "Report Version: 3.0"
            )
        if "Next Review: February 10, 2026" in para.text:
            para.text = para.text.replace(
                "Next Review: February 10, 2026",
                "Next Review: February 17, 2026"
            )

    # Update executive summary
    for para in doc.paragraphs:
        if "The TTG Distributed Computation System successfully processes" in para.text:
            para.text = (
                "The TTG Distributed Computation System processes large parameter sets across "
                "multiple Kubernetes worker nodes. The project has successfully completed three milestones "
                "and now supports dual queue backends (Redis Streams and RabbitMQ) with verified "
                "fault tolerance, automatic retry/dead-letter handling, and visual monitoring."
            )
            break


def main():
    print(f"Loading template: {TEMPLATE_PATH}")
    if not os.path.exists(TEMPLATE_PATH):
        print(f"ERROR: Template not found at {TEMPLATE_PATH}")
        sys.exit(1)

    doc = Document(TEMPLATE_PATH)

    print("Updating report metadata...")
    update_report_metadata(doc)

    print("Updating milestone overview table...")
    update_milestone_table(doc)

    # Find insertion point: just before "Architecture Evolution" heading
    # We insert M3 section between M2 content and Architecture Evolution
    insert_before_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "Architecture Evolution" in para.text:
            insert_before_idx = i
            break

    # Since python-docx doesn't support inserting paragraphs at arbitrary positions
    # easily, we'll append M3 sections at the end of the document (before Next Steps)
    # The approach: add M3 content, then architecture comparison

    print("Adding Milestone 3 section...")
    add_m3_section(doc)

    print("Adding backend comparison...")
    add_backend_comparison(doc)

    print("Adding M3 key files...")
    add_m3_files(doc)

    print("Updating comparison table...")
    update_comparison_table(doc)

    print("Updating next steps...")
    update_next_steps(doc)

    print("Updating demo commands...")
    update_demo_commands(doc)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Report saved: {OUTPUT_PATH}")
    print(f"   File size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    main()
