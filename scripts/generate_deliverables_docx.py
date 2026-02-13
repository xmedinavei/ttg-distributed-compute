#!/usr/bin/env python3
"""
Generate DOCX versions of all AME-UP deliverable documents.

Converts the Markdown deliverables to clean, formal Word documents
with consistent styling matching the existing TTG report format.

Usage:
    python3 scripts/generate_deliverables_docx.py

Output:
    deliverables/docx/D1_Optimization_Performance_Report.docx
    deliverables/docx/D2_Technical_Artifacts_Index.docx
    deliverables/docx/D2_Operational_Runbook.docx
    deliverables/docx/D3_Final_Presentation.docx
    deliverables/docx/D3_Engineering_Playbook.docx
"""

import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx")
    sys.exit(1)

PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DELIVERABLES_DIR = os.path.join(PROJECT_DIR, "deliverables")
OUTPUT_DIR = os.path.join(DELIVERABLES_DIR, "docx")

# Deliverable files to convert
DELIVERABLES = [
    "D1_Optimization_Performance_Report.md",
    "D2_Technical_Artifacts_Index.md",
    "D2_Operational_Runbook.md",
    "D3_Final_Presentation.md",
    "D3_Engineering_Playbook.md",
]


# ═══════════════════════════════════════════════════════════════════════════
# STYLE HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def set_font(run, name="Times New Roman", size=11, bold=False,
             underline=False, italic=False, color=None):
    """Configure font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    size_map = {1: 14, 2: 13, 3: 12}
    size = size_map.get(level, 11)

    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, underline=(level <= 2))
    return p


def add_body(doc, text, size=11):
    """Add a justified body paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_code_block(doc, text):
    """Add a code block as a formatted paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=8)
    return p


def add_bullet(doc, text, size=10):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("  " + text)
    set_font(run, size=size)
    p.style = doc.styles["List Bullet"]
    return p


def make_table(doc, headers, rows):
    """Create a formatted table."""
    if not headers or not rows:
        return None

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.strip())
        set_font(run, size=9, bold=True)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(min(len(row_data), num_cols)):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            cell_text = row_data[c_idx].strip() if c_idx < len(row_data) else ""
            run = p.add_run(cell_text)
            is_bold = (c_idx == 0)
            set_font(run, size=9, bold=is_bold)

    return table


# ═══════════════════════════════════════════════════════════════════════════
# MARKDOWN PARSER
# ═══════════════════════════════════════════════════════════════════════════

def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (headers, rows, end_idx)."""
    headers = []
    rows = []
    idx = start_idx

    # Parse header row
    header_line = lines[idx].strip()
    if "|" in header_line:
        headers = [c.strip() for c in header_line.split("|") if c.strip()]
        idx += 1

    # Skip separator row (|---|---|...)
    if idx < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|", lines[idx]):
        idx += 1

    # Parse data rows
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row_line = lines[idx].strip()
        cells = [c.strip() for c in row_line.split("|") if c.strip()]
        if cells:
            rows.append(cells)
        idx += 1

    return headers, rows, idx


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a DOCX document."""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    idx = 0
    in_code_block = False
    code_lines = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Skip front matter lines (---, metadata lines at top)
        if stripped == "---":
            idx += 1
            continue

        # Skip HTML comments
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            idx += 1
            continue

        # Code block toggle
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                if code_lines:
                    add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            add_heading(doc, stripped[2:], level=1)
            idx += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], level=2)
            idx += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], level=3)
            idx += 1
            continue

        # Tables
        if "|" in stripped and not stripped.startswith(">"):
            # Check if next line is a separator (confirming it's a table)
            if (idx + 1 < len(lines) and
                    re.match(r"^\s*\|[\s\-:|]+\|", lines[idx + 1])):
                headers, rows, new_idx = parse_table(lines, idx)
                if headers and rows:
                    make_table(doc, headers, rows)
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                idx = new_idx
                continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            # Clean markdown formatting
            bullet_text = re.sub(r"\*\*(.*?)\*\*", r"\1", bullet_text)
            bullet_text = re.sub(r"`(.*?)`", r"\1", bullet_text)
            add_bullet(doc, bullet_text)
            idx += 1
            continue

        # Numbered items
        if re.match(r"^\d+\.\s", stripped):
            item_text = re.sub(r"^\d+\.\s+", "", stripped)
            item_text = re.sub(r"\*\*(.*?)\*\*", r"\1", item_text)
            item_text = re.sub(r"`(.*?)`", r"\1", item_text)
            add_bullet(doc, item_text)
            idx += 1
            continue

        # Block quotes
        if stripped.startswith("> "):
            quote_text = stripped[2:]
            quote_text = re.sub(r"\*\*(.*?)\*\*", r"\1", quote_text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(quote_text)
            set_font(run, size=10, italic=True)
            idx += 1
            continue

        # Regular text (skip empty lines, metadata lines)
        if stripped and not stripped.startswith("**Date:") and \
           not stripped.startswith("**Project:") and \
           not stripped.startswith("**Version:") and \
           not stripped.startswith("**Author:") and \
           not stripped.startswith("**Program:"):
            # Clean markdown formatting
            text = stripped
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            text = re.sub(r"`(.*?)`", r"\1", text)
            text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)

            if text:
                add_body(doc, text)

        idx += 1

    # Save
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    return os.path.getsize(docx_path)


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 70)
    print("  TTG Deliverables DOCX Generator")
    print("=" * 70)
    print()

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    total_files = 0
    total_size = 0

    for md_file in DELIVERABLES:
        md_path = os.path.join(DELIVERABLES_DIR, md_file)
        docx_file = md_file.replace(".md", ".docx")
        docx_path = os.path.join(OUTPUT_DIR, docx_file)

        if not os.path.exists(md_path):
            print(f"  SKIP: {md_file} (not found)")
            continue

        print(f"  Converting: {md_file}")
        size = convert_md_to_docx(md_path, docx_path)
        print(f"       Saved: {docx_path}")
        print(f"       Size:  {size:,} bytes")
        print()
        total_files += 1
        total_size += size

    print("-" * 70)
    print(f"  Done! Generated {total_files} DOCX files ({total_size:,} bytes total)")
    print(f"  Output directory: {OUTPUT_DIR}")
    print()


if __name__ == "__main__":
    main()
